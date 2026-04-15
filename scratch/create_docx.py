from docx import Document
import os

def create_sample_docx(file_path):
    doc = Document()
    doc.add_heading('SmartDoc AI Test Document', 0)
    
    doc.add_heading('Section 1: Introduction', level=1)
    doc.add_paragraph('This is a test document for SmartDoc AI Phase 1.')
    doc.add_paragraph('It contains some text to verify the DOCX loader functionality.')
    
    doc.add_heading('Section 2: Technical Details', level=1)
    doc.add_paragraph('SmartDoc AI uses RAG (Retrieval-Augmented Generation) to answer questions.')
    doc.add_paragraph('We are testing the Docx2txtLoader integration.')
    
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    doc.save(file_path)
    print(f"Sample DOCX created at: {file_path}")

if __name__ == "__main__":
    path = r"e:\VSCode\PhatTrienPhanMemMaNguonMo\SmartDoc-AI\data\samples\sample_test.docx"
    create_sample_docx(path)
